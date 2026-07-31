"""
Ingestion for the single user input: a dump of everything about them.

Accepts LaTeX, Markdown, plain text, Word, or PDF. Text-like formats are
decoded locally; PDFs are passed to Claude natively as document blocks, so
they are returned as base64 rather than extracted here.
"""

import base64
from dataclasses import dataclass
from io import BytesIO
from typing import Optional

TEXT_EXTENSIONS = {".tex", ".md", ".markdown", ".txt", ".text", ".rtf"}


@dataclass
class Dump:
    """The candidate's ingested information dump."""
    text: Optional[str] = None
    pdf_base64: Optional[str] = None
    filename: str = ""

    def as_content_blocks(self, prompt: str) -> list:
        """Build the user-content blocks for a Claude message carrying this dump."""
        blocks = []
        if self.pdf_base64:
            blocks.append({
                "type": "document",
                "source": {"type": "base64", "media_type": "application/pdf", "data": self.pdf_base64},
            })
        if self.text:
            blocks.append({"type": "text", "text": f"CANDIDATE INFORMATION DUMP:\n\n{self.text}"})
        blocks.append({"type": "text", "text": prompt})
        return blocks


def ingest_file(filename: str, data: bytes) -> Dump:
    """Turn an uploaded file into a Dump. Raises ValueError on unsupported input."""
    if not data or not data.strip():
        raise ValueError("That file is empty — upload one that actually contains your information.")

    ext = ("." + filename.rsplit(".", 1)[-1].lower()) if "." in filename else ""

    if ext == ".pdf" or data[:5] == b"%PDF-":
        if data[:5] != b"%PDF-":
            raise ValueError("That .pdf file isn't a valid PDF. Re-export it and try again.")
        return Dump(pdf_base64=base64.standard_b64encode(data).decode(), filename=filename)

    if ext == ".docx":
        return Dump(text=_extract_docx(data), filename=filename)

    if ext == ".doc":
        raise ValueError("Legacy .doc files are not supported — save as .docx or PDF and try again.")

    if ext in TEXT_EXTENSIONS or not ext:
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError:
            text = data.decode("latin-1")
        if not text.strip():
            raise ValueError("That file contains no readable text.")
        return Dump(text=text, filename=filename)

    raise ValueError(f"Unsupported file type '{ext}'. Upload LaTeX, Markdown, plain text, Word (.docx), or PDF.")


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    if not text.strip():
        raise ValueError("Could not extract any text from the Word document.")
    return text
